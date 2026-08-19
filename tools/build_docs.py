"""Turn START-HERE.md into a PDF and a Word file for the client.

The markdown is the source of truth and stays that way -- both outputs are
generated, never hand-edited, so a correction lands in one place and reaches
all three.

PDF via Chromium's own print engine, which is already installed for Stage B
rendering. It is a real typesetter: proper hyphenation-free justification,
widow control, and page breaks we can steer with CSS. A hand-rolled PDF library
would take longer and read worse.

DOCX via python-docx, because a client who wants Word usually wants to EDIT it
-- adding their own contact details, cutting a section for a colleague. A PDF
exported to .docx is a picture of a document; this is a document.

Usage:  python tools/build_docs.py [--source START-HERE.md] [--out-dir .]
"""

from __future__ import annotations

import argparse
import pathlib
import re
import sys

ROOT = pathlib.Path(__file__).resolve().parents[1]

# Print CSS. Kept here rather than in a file because it is meaningless on its
# own and would rot in a directory of its own.
STYLE = """
@page { size: A4; margin: 18mm 16mm 20mm 16mm; }

:root {
  --ink:     #14171f;
  --muted:   #5b6472;
  --rule:    #d9dee6;
  --accent:  #1c3f6e;
  --wash:    #f4f6f9;
}

* { box-sizing: border-box; }

body {
  font-family: "Segoe UI", -apple-system, "Helvetica Neue", Arial, sans-serif;
  font-size: 10.5pt;
  line-height: 1.55;
  color: var(--ink);
  margin: 0;
}

h1 {
  font-size: 24pt;
  line-height: 1.15;
  margin: 0 0 4pt;
  color: var(--accent);
  letter-spacing: -0.01em;
}

h2 {
  font-size: 14pt;
  margin: 22pt 0 6pt;
  padding-top: 8pt;
  border-top: 1.5px solid var(--rule);
  color: var(--accent);
  /* A heading alone at the foot of a page is the commonest ugliness in a
     generated PDF, and the one readers actually notice. */
  break-after: avoid;
  break-inside: avoid;
}

h3 {
  font-size: 11.5pt;
  margin: 14pt 0 4pt;
  color: var(--ink);
  break-after: avoid;
}

p, li { orphans: 3; widows: 3; }
p { margin: 0 0 8pt; }
ul, ol { margin: 0 0 8pt; padding-left: 18pt; }
li { margin: 0 0 3pt; }

a { color: var(--accent); text-decoration: none; }

code {
  font-family: "Cascadia Mono", Consolas, "SF Mono", monospace;
  font-size: 9pt;
  background: var(--wash);
  border: 1px solid var(--rule);
  border-radius: 2px;
  padding: 0.5pt 3pt;
}

pre {
  background: var(--wash);
  border: 1px solid var(--rule);
  border-left: 3px solid var(--accent);
  border-radius: 2px;
  padding: 8pt 10pt;
  margin: 0 0 10pt;
  overflow: visible;
  white-space: pre-wrap;
  word-break: break-word;
  break-inside: avoid;
}

pre code { background: none; border: 0; padding: 0; font-size: 8.5pt; }

table {
  width: 100%;
  border-collapse: collapse;
  margin: 0 0 10pt;
  font-size: 9.5pt;
  break-inside: avoid;
}

th {
  text-align: left;
  font-weight: 600;
  background: var(--wash);
  border-bottom: 1.5px solid var(--rule);
  padding: 5pt 7pt;
}

td {
  border-bottom: 1px solid var(--rule);
  padding: 5pt 7pt;
  vertical-align: top;
}

/* The markdown uses an empty header row for two-column definition tables;
   an empty band of grey looks like a mistake. */
thead tr:has(th:empty:only-of-type),
thead:has(th:empty) { display: none; }

/* The source uses `---` between sections AND every h2 draws its own rule, so
   rendering both gives a doubled line at every heading. The heading's rule is
   the one tied to the thing it separates, so it wins. */
hr { display: none; }

blockquote {
  margin: 0 0 10pt;
  padding: 6pt 10pt;
  border-left: 3px solid var(--accent);
  background: var(--wash);
  color: var(--muted);
}

.subtitle {
  color: var(--muted);
  font-size: 10pt;
  margin: 0 0 14pt;
  padding-bottom: 10pt;
  border-bottom: 1.5px solid var(--rule);
}
"""


def to_html(markdown_text: str, title: str, subtitle: str) -> str:
    import markdown as md

    body = md.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "sane_lists", "attr_list"],
    )
    # The document's own H1 is replaced by the styled header below it, so the
    # title is not printed twice.
    body = re.sub(r"<h1>.*?</h1>", "", body, count=1, flags=re.S)
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{title}</title><style>{STYLE}</style></head><body>"
        f"<h1>{title}</h1><p class='subtitle'>{subtitle}</p>"
        f"{body}</body></html>"
    )


def write_pdf(html: str, out: pathlib.Path) -> None:
    from playwright.sync_api import sync_playwright

    tmp = out.with_suffix(".tmp.html")
    tmp.write_text(html, encoding="utf-8")
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(tmp.resolve().as_uri(), wait_until="networkidle")
            page.pdf(
                path=str(out),
                format="A4",
                print_background=True,
                display_header_footer=True,
                header_template="<div></div>",
                footer_template=(
                    "<div style='width:100%;font-size:7.5pt;color:#8a929e;"
                    "font-family:Segoe UI,Arial,sans-serif;padding:0 16mm;"
                    "display:flex;justify-content:space-between'>"
                    "<span>haat-lister — start here</span>"
                    "<span class='pageNumber'></span></div>"
                ),
                margin={"top": "16mm", "bottom": "18mm", "left": "16mm", "right": "16mm"},
            )
            browser.close()
    finally:
        tmp.unlink(missing_ok=True)


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def _runs(paragraph, text: str) -> None:
    """Bold, italic and code, applied inline. Everything else is plain."""
    for piece in _INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
        elif piece.startswith("*") and piece.endswith("*"):
            paragraph.add_run(piece[1:-1]).italic = True
        else:
            paragraph.add_run(piece)


def write_docx(markdown_text: str, out: pathlib.Path, title: str, subtitle: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    document.add_heading(title, level=0)
    lead = document.add_paragraph(subtitle)
    lead.runs[0].italic = True
    lead.runs[0].font.color.rgb = RGBColor(0x5B, 0x64, 0x72)

    lines = markdown_text.splitlines()
    i = 0
    seen_title = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code
        if stripped.startswith("```"):
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].strip().replace("|", "").replace(" ", "")
        ) <= {"-", ":"}:
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = document.add_table(rows=0, cols=len(header))
            table.style = "Light Grid Accent 1"
            if any(header):
                cells = table.add_row().cells
                for cell, text in zip(cells, header, strict=False):
                    cell.text = ""
                    _runs(cell.paragraphs[0], text)
            for row in rows:
                cells = table.add_row().cells
                for cell, text in zip(cells, row, strict=False):
                    cell.text = ""
                    _runs(cell.paragraphs[0], text)
            document.add_paragraph()
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            if level == 1 and not seen_title:
                seen_title = True  # already the document title
                i += 1
                continue
            document.add_heading(re.sub(r"[`*]", "", text), level=min(level, 4))
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            _runs(document.add_paragraph(style="List Bullet"), stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            _runs(
                document.add_paragraph(style="List Number"),
                re.sub(r"^\d+\.\s", "", stripped),
            )
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # A paragraph is every following non-blank, non-structural line.
        buffer = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", "```", "- ", "* ", "---")):
                break
            if re.match(r"^\d+\.\s", nxt):
                break
            buffer.append(nxt)
            i += 1
        _runs(document.add_paragraph(), " ".join(buffer))

    document.save(out)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=pathlib.Path, default=ROOT / "START-HERE.md")
    parser.add_argument("--out-dir", type=pathlib.Path, default=ROOT)
    args = parser.parse_args()

    text = args.source.read_text(encoding="utf-8")
    title = "haat-lister — start here"
    subtitle = (
        "Product page URLs into a haat bulk-listing CSV. "
        "Install, first run, and the three decisions the tool leaves to you."
    )

    args.out_dir.mkdir(parents=True, exist_ok=True)
    pdf = args.out_dir / "START-HERE.pdf"
    docx = args.out_dir / "START-HERE.docx"

    write_pdf(to_html(text, title, subtitle), pdf)
    write_docx(text, docx, title, subtitle)

    for path in (pdf, docx):
        print(f"  {path.name:20} {path.stat().st_size / 1024:6.0f} KB")
    return 0


if __name__ == "__main__":
    sys.exit(main())
